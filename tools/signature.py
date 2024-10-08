from docx import Document
from docx.shared import Inches

def center_insert_img(doc, ZDLS,XZZZ,XZCY):
    # print("插入签名")
    have_ZDLS=False
    have_XZZZ=False
    have_XZCY=False
    is_zq = False
    tables = doc.tables
    table = tables[0]
    for i in range(0, len(table.rows)):
        for j in range(0, len(table.columns)):
            for p in table.cell(i, j).paragraphs:
                p.paragraph_format.line_spacing = 1
                # print(p.text)
                if ('导师签字：' in p.text or '指导老师签字：' in p.text or '指导教师：'in p.text or '指导老师（签字）：' in p.text or '指导教师（签字）：' in p.text or '是否符合预定计划：                        指导老师(签字)：' in p.text ) and have_ZDLS==False:
                    # print('找到了导师签字')
                    run = p.add_run()
                    # run.add_break()
                    # 添加图片并指定大小
                    for img in ZDLS:
                        # print("导师签字")
                        run.add_picture(img, height=Inches(0.35))
                    have_ZDLS = True

                if '小组组长（签字）：' in p.text or '答辩组组长（签字）：' in p.text and  have_XZZZ==False:
                    # print('开题检查评审小组组长（签字）')
                    run = p.add_run()
                    # run.add_break()
                    # 添加图片并指定大小
                    for img in XZZZ:
                        run.add_picture(img, height=Inches(0.35))
                    have_XZZZ = True

                if '小组成员（签字）：' in p.text and   have_XZCY==False:
                    # print('找到了开题检查评审小组成员（签字）：')
                    run = p.add_run()
                    # run.add_break()
                    # 添加图片并指定大小
                    for img in XZCY:
                        run.add_picture(img, height=Inches(0.5))
                    have_XZCY = True
                if '中期结果' in p.text and is_zq == False:
                    is_zq = True
                if '指导老师(签字)：' in p.text and is_zq == True and have_ZDLS == False:
                    run = p.add_run()
                    # run.add_break()
                    # 添加图片并指定大小
                    for img in ZDLS:
                        run.add_picture(img, height=Inches(0.35))
                        run.add_break()
                    # run.add_text('                                                       年    月    日')
                    have_ZDLS = True

def signature_fill(docx_path, ZDLS,XZZZ,XZCY):
    document = Document(docx_path)
    # 插入图片居中
    center_insert_img(document, ZDLS, XZZZ, XZCY)
    # 保存结果文件
    # print('执行到此')
    document.save(docx_path)



